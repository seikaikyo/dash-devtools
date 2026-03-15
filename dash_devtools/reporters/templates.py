"""
報告模板與樣式模組

提供 Word 報告用的：
- 儲存格樣式設定
- 時間格式化
- 測試類型標籤與說明
- 色彩常數
- 測試案例文件建構
"""

from pathlib import Path
from typing import Dict, Tuple

try:
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ========== 色彩常數 ==========

# 狀態色
COLOR_PASS_RGB = (76, 175, 80)       # 綠色
COLOR_FAIL_RGB = (244, 67, 54)       # 紅色
COLOR_SKIP_RGB = (255, 193, 7)       # 黃色
COLOR_GREY_RGB = (128, 128, 128)     # 灰色
COLOR_LIGHT_GREY_RGB = (150, 150, 150)
COLOR_SUBTITLE_RGB = (100, 100, 100)

# 表格色
COLOR_HEADER_HEX = '2196F3'          # 藍色表頭
COLOR_LABEL_BG_HEX = 'F5F5F5'        # 標籤背景
COLOR_PASS_BG_HEX = 'C8E6C9'         # 淺綠背景
COLOR_FAIL_BG_HEX = 'FFCDD2'         # 淺紅背景

# API / 終端色
COLOR_API_RGB = (33, 150, 243)        # 藍色
COLOR_TERMINAL_RGB = (156, 39, 176)   # 紫色
COLOR_CODE_RGB = (80, 80, 80)         # 深灰


# ========== 測試類型對照 ==========

TYPE_LABELS = {
    'UIT': '單元測試 (UIT)',
    'SMOKE': '煙霧測試 (Smoke)',
    'E2E': '端對端測試 (E2E)',
    'UAT': '驗收測試 (UAT)'
}

TYPE_DESCRIPTIONS = [
    ('UIT (Unit Integration Testing)',
     '單元測試驗證各個模組、函數的正確性。使用 Vitest/Jest 框架執行，並產生程式碼覆蓋率報告。'),
    ('Smoke Test (煙霧測試)',
     '快速驗證系統關鍵路徑是否正常運作。包含應用程式啟動、頁面載入、API 健康檢查等基本功能。'),
    ('E2E (End-to-End Testing)',
     '端對端測試模擬真實使用情境，驗證完整的使用者流程。使用 Playwright 自動化測試框架執行。'),
    ('UAT (User Acceptance Testing)',
     '使用者驗收測試從業務角度驗證系統符合需求規格。測試案例依據使用者角色設計，確保系統滿足業務需求。'),
]


# ========== 樣式輔助函數 ==========

def set_cell_shading(cell, color: str):
    """設定表格儲存格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def format_duration(duration: float) -> str:
    """
    智慧格式化時間

    Args:
        duration: 秒數

    Returns:
        格式化後的字串 (如 '1.2s', '150ms', '500us', '-')
    """
    if duration <= 0:
        return '-'
    elif duration < 0.001:  # < 1ms
        return f'{duration * 1000000:.0f}us'
    elif duration < 0.1:  # < 100ms
        return f'{duration * 1000:.2f}ms'
    elif duration < 1:  # < 1s
        return f'{duration * 1000:.0f}ms'
    else:
        return f'{duration:.1f}s'


def format_duration_detail(duration: float) -> str:
    """
    智慧格式化時間 (較精確版本，用於測試案例明細)

    Args:
        duration: 秒數

    Returns:
        格式化後的字串
    """
    if not duration or duration <= 0:
        return ''
    elif duration < 0.001:  # < 1ms
        return f'{duration * 1000000:.0f}us'
    elif duration < 0.1:  # < 100ms
        return f'{duration * 1000:.2f}ms'
    elif duration < 1:  # < 1s
        return f'{duration * 1000:.0f}ms'
    else:
        return f'{duration:.2f}s'


def format_table_duration(duration: float) -> str:
    """
    格式化表格中的時間欄位

    Args:
        duration: 秒數

    Returns:
        格式化後的字串
    """
    if duration <= 0:
        return '-'
    elif duration < 0.1:  # < 100ms
        return f'{duration * 1000:.2f}ms'
    elif duration < 1:  # < 1s
        return f'{duration * 1000:.0f}ms'
    else:
        return f'{duration:.1f}s'


def rgb(color_tuple: Tuple[int, int, int]) -> 'RGBColor':
    """從 tuple 建立 RGBColor"""
    return RGBColor(*color_tuple)


# ========== 文件建構輔助函數 ==========

def build_single_test_case(doc, tc: Dict, index: int):
    """建立單一測試案例區塊

    Args:
        doc: python-docx Document 物件
        tc: 測試案例字典
        index: 案例序號
    """
    test_name = tc.get('name', '')
    status = tc.get('status', 'passed')
    duration = tc.get('duration', 0)
    screenshot_path = tc.get('screenshot', '')

    # 測試案例標題
    p = doc.add_paragraph()

    # 狀態標示
    status_map = {
        'passed': ('[PASS] ', COLOR_PASS_RGB),
        'failed': ('[FAIL] ', COLOR_FAIL_RGB),
    }
    label, color = status_map.get(status, ('[SKIP] ', COLOR_SKIP_RGB))
    status_run = p.add_run(label)
    status_run.font.color.rgb = rgb(color)
    status_run.bold = True

    # 測試名稱
    name_run = p.add_run(f'{index}. {test_name}')
    name_run.font.size = Pt(11)

    # 時間
    time_str = format_duration_detail(duration)
    if time_str:
        time_run = p.add_run(f'  ({time_str})')
        time_run.font.size = Pt(9)
        time_run.font.color.rgb = rgb(COLOR_GREY_RGB)

    # 截圖
    if screenshot_path and Path(screenshot_path).exists():
        doc.add_picture(screenshot_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        build_text_output(doc, tc)

    # 錯誤訊息
    error = tc.get('error', '')
    if error:
        error_p = doc.add_paragraph()
        error_run = error_p.add_run(f'Error: {error[:300]}')
        error_run.font.size = Pt(9)
        error_run.font.color.rgb = rgb(COLOR_FAIL_RGB)

    doc.add_paragraph()  # 間隔


def build_text_output(doc, tc: Dict):
    """建立 API 回應或終端輸出區塊

    Args:
        doc: python-docx Document 物件
        tc: 測試案例字典
    """
    api_response = tc.get('api_response', '')
    terminal_output = tc.get('terminal_output', '')

    if api_response:
        api_label = doc.add_paragraph()
        api_label_run = api_label.add_run('API Response:')
        api_label_run.font.size = Pt(10)
        api_label_run.font.color.rgb = rgb(COLOR_API_RGB)
        api_label_run.bold = True

        api_para = doc.add_paragraph()
        api_para.paragraph_format.left_indent = Inches(0.3)
        display_response = api_response[:500] + ('...' if len(api_response) > 500 else '')
        api_run = api_para.add_run(display_response)
        api_run.font.size = Pt(9)
        api_run.font.name = 'Consolas'
        api_run.font.color.rgb = rgb(COLOR_CODE_RGB)

    elif terminal_output:
        term_label = doc.add_paragraph()
        term_label_run = term_label.add_run('Terminal Output:')
        term_label_run.font.size = Pt(10)
        term_label_run.font.color.rgb = rgb(COLOR_TERMINAL_RGB)
        term_label_run.bold = True

        term_para = doc.add_paragraph()
        term_para.paragraph_format.left_indent = Inches(0.3)
        display_output = terminal_output[:400] + ('...' if len(terminal_output) > 400 else '')
        term_run = term_para.add_run(display_output)
        term_run.font.size = Pt(9)
        term_run.font.name = 'Consolas'
        term_run.font.color.rgb = rgb(COLOR_CODE_RGB)
