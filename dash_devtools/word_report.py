"""
Word 測試報告生成模組

生成包含圖文內容的專業測試報告
- 封面頁
- 測試摘要
- 測試結果表格
- 圖表 (通過率圓餅圖、各類型長條圖)
- 截圖 (可選)
"""

import io
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use('Agg')  # 非互動模式
    import matplotlib.pyplot as plt
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

from rich.console import Console

console = Console()


def set_cell_shading(cell, color: str):
    """設定表格儲存格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_pass_rate_chart(passed: int, failed: int) -> Optional[bytes]:
    """建立通過率圓餅圖"""
    if not HAS_MATPLOTLIB:
        return None

    # 設定中文字型 (macOS)
    plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'Heiti TC', 'PingFang TC', 'STHeiti']
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(4, 4))

    if passed + failed == 0:
        sizes = [1]
        colors = ['#E0E0E0']
        labels = ['無測試']
    else:
        sizes = [passed, failed] if failed > 0 else [passed]
        colors = ['#4CAF50', '#F44336'] if failed > 0 else ['#4CAF50']
        labels = ['通過', '失敗'] if failed > 0 else ['通過']

    wedges, texts, autotexts = ax.pie(
        sizes,
        labels=labels,
        colors=colors,
        autopct='%1.1f%%',
        startangle=90,
        textprops={'fontsize': 12}
    )

    ax.set_title('測試通過率', fontsize=14, fontweight='bold')

    # 儲存為 bytes
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def create_test_type_chart(results: Dict) -> Optional[bytes]:
    """建立各類型測試長條圖"""
    if not HAS_MATPLOTLIB:
        return None

    # 設定中文字型
    plt.rcParams['font.sans-serif'] = ['Arial Unicode MS', 'Heiti TC', 'PingFang TC', 'Microsoft JhengHei']
    plt.rcParams['axes.unicode_minus'] = False

    fig, ax = plt.subplots(figsize=(8, 4))

    types = list(results.keys())
    passed = [results[t].get('passed', 0) for t in types]
    failed = [results[t].get('failed', 0) for t in types]

    x = range(len(types))
    width = 0.35

    bars1 = ax.bar([i - width/2 for i in x], passed, width, label='通過', color='#4CAF50')
    bars2 = ax.bar([i + width/2 for i in x], failed, width, label='失敗', color='#F44336')

    ax.set_ylabel('測試數量')
    ax.set_title('各類型測試結果', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(types)
    ax.legend()

    # 加上數值標籤
    for bar in bars1:
        height = bar.get_height()
        if height > 0:
            ax.annotate(f'{int(height)}',
                       xy=(bar.get_x() + bar.get_width() / 2, height),
                       xytext=(0, 3),
                       textcoords="offset points",
                       ha='center', va='bottom', fontsize=10)

    for bar in bars2:
        height = bar.get_height()
        if height > 0:
            ax.annotate(f'{int(height)}',
                       xy=(bar.get_x() + bar.get_width() / 2, height),
                       xytext=(0, 3),
                       textcoords="offset points",
                       ha='center', va='bottom', fontsize=10)

    plt.tight_layout()

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def generate_word_report(
    project_name: str,
    test_results: Dict,
    output_path: str,
    screenshots: List[str] = None,
    include_charts: bool = True
) -> str:
    """
    生成 Word 測試報告

    Args:
        project_name: 專案名稱
        test_results: 測試結果字典 (來自 test_suite)
        output_path: 輸出檔案路徑
        screenshots: 截圖路徑列表 (可選)
        include_charts: 是否包含圖表

    Returns:
        輸出檔案路徑
    """
    if not HAS_DOCX:
        raise ImportError("請安裝 python-docx: pip install python-docx")

    doc = Document()

    # 設定文件樣式
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(12)

    # ========== 封面頁 ==========
    # 標題
    title = doc.add_heading('', level=0)
    title_run = title.add_run(f'{project_name} 測試報告')
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副標題
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run('四大類型測試套件執行結果')
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    # 日期
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp = test_results.get('timestamp', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
    date_run = date_para.add_run(f'報告產生時間: {timestamp}')
    date_run.font.size = Pt(14)

    # 狀態標示
    doc.add_paragraph()
    doc.add_paragraph()
    status_para = doc.add_paragraph()
    status_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    overall_success = test_results.get('overall_success', True)
    if overall_success:
        status_run = status_para.add_run('ALL TESTS PASSED')
        status_run.font.size = Pt(24)
        status_run.font.bold = True
        status_run.font.color.rgb = RGBColor(76, 175, 80)  # 綠色
    else:
        status_run = status_para.add_run('SOME TESTS FAILED')
        status_run.font.size = Pt(24)
        status_run.font.bold = True
        status_run.font.color.rgb = RGBColor(244, 67, 54)  # 紅色

    doc.add_page_break()

    # ========== 測試摘要 ==========
    doc.add_heading('測試摘要', level=1)

    summary = test_results.get('summary', {})
    total_passed = summary.get('total_passed', 0)
    total_failed = summary.get('total_failed', 0)
    total_duration = summary.get('total_duration', 0)
    coverage = summary.get('coverage', 0)

    # 摘要表格
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'

    summary_data = [
        ('總測試數', str(total_passed + total_failed)),
        ('通過', str(total_passed)),
        ('失敗', str(total_failed)),
        ('執行時間', f'{total_duration:.1f} 秒'),
        ('程式碼覆蓋率', f'{coverage:.1f}%' if coverage > 0 else 'N/A'),
    ]

    for i, (label, value) in enumerate(summary_data):
        row = summary_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # 設定標籤欄背景色
        set_cell_shading(row.cells[0], 'F5F5F5')

    doc.add_paragraph()

    # ========== 通過率圖表 ==========
    if include_charts and HAS_MATPLOTLIB:
        doc.add_heading('測試通過率', level=2)

        chart_data = create_pass_rate_chart(total_passed, total_failed)
        if chart_data:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(chart_data)
                tmp_path = tmp.name

            doc.add_picture(tmp_path, width=Inches(3.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            Path(tmp_path).unlink(missing_ok=True)

        doc.add_paragraph()

    # ========== 各類型測試結果 ==========
    doc.add_heading('各類型測試結果', level=1)

    tests = test_results.get('tests', {})

    # 結果表格
    result_table = doc.add_table(rows=len(tests) + 1, cols=5)
    result_table.style = 'Table Grid'

    # 表頭
    header_row = result_table.rows[0]
    headers = ['測試類型', '狀態', '通過', '失敗', '時間']
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        set_cell_shading(cell, '2196F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 資料列
    type_labels = {
        'UIT': '單元測試 (UIT)',
        'SMOKE': '煙霧測試 (Smoke)',
        'E2E': '端對端測試 (E2E)',
        'UAT': '驗收測試 (UAT)'
    }

    for i, (test_type, result) in enumerate(tests.items(), start=1):
        row = result_table.rows[i]
        row.cells[0].text = type_labels.get(test_type, test_type)

        success = result.get('success', True)
        row.cells[1].text = 'PASS' if success else 'FAIL'
        # 狀態顏色
        status_cell = row.cells[1]
        if success:
            set_cell_shading(status_cell, 'C8E6C9')  # 淺綠
        else:
            set_cell_shading(status_cell, 'FFCDD2')  # 淺紅

        row.cells[2].text = str(result.get('passed', 0))
        row.cells[3].text = str(result.get('failed', 0))
        row.cells[4].text = f"{result.get('duration', 0):.1f}s"

    doc.add_paragraph()

    # ========== 詳細測試案例列表 (含截圖) ==========
    for test_type, result in tests.items():
        test_cases = result.get('test_cases', [])
        if not test_cases:
            continue

        doc.add_page_break()
        doc.add_heading(f'{type_labels.get(test_type, test_type)} - 測試案例明細', level=1)

        for i, tc in enumerate(test_cases, start=1):
            test_name = tc.get('name', '')
            status = tc.get('status', 'passed')
            duration = tc.get('duration', 0)
            screenshot_path = tc.get('screenshot', '')

            # 測試案例標題
            p = doc.add_paragraph()
            # 狀態圖示
            if status == 'passed':
                status_run = p.add_run('[PASS] ')
                status_run.font.color.rgb = RGBColor(76, 175, 80)
                status_run.bold = True
            elif status == 'failed':
                status_run = p.add_run('[FAIL] ')
                status_run.font.color.rgb = RGBColor(244, 67, 54)
                status_run.bold = True
            else:
                status_run = p.add_run('[SKIP] ')
                status_run.font.color.rgb = RGBColor(255, 193, 7)
                status_run.bold = True

            # 測試名稱
            name_run = p.add_run(f'{i}. {test_name}')
            name_run.font.size = Pt(11)

            # 時間
            if duration:
                time_run = p.add_run(f'  ({duration:.2f}s)')
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(128, 128, 128)

            # 截圖
            if screenshot_path and Path(screenshot_path).exists():
                doc.add_picture(screenshot_path, width=Inches(5.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # API 回應 (當沒有截圖時顯示)
                api_response = tc.get('api_response', '')
                terminal_output = tc.get('terminal_output', '')

                if api_response:
                    api_label = doc.add_paragraph()
                    api_label_run = api_label.add_run('API Response:')
                    api_label_run.font.size = Pt(10)
                    api_label_run.font.color.rgb = RGBColor(33, 150, 243)
                    api_label_run.bold = True

                    # 格式化 JSON 顯示
                    api_para = doc.add_paragraph()
                    api_para.paragraph_format.left_indent = Inches(0.3)
                    # 限制顯示長度
                    display_response = api_response[:500] + ('...' if len(api_response) > 500 else '')
                    api_run = api_para.add_run(display_response)
                    api_run.font.size = Pt(9)
                    api_run.font.name = 'Consolas'
                    api_run.font.color.rgb = RGBColor(80, 80, 80)

                elif terminal_output:
                    # 終端輸出 (UIT 測試)
                    term_label = doc.add_paragraph()
                    term_label_run = term_label.add_run('Terminal Output:')
                    term_label_run.font.size = Pt(10)
                    term_label_run.font.color.rgb = RGBColor(156, 39, 176)
                    term_label_run.bold = True

                    term_para = doc.add_paragraph()
                    term_para.paragraph_format.left_indent = Inches(0.3)
                    display_output = terminal_output[:400] + ('...' if len(terminal_output) > 400 else '')
                    term_run = term_para.add_run(display_output)
                    term_run.font.size = Pt(9)
                    term_run.font.name = 'Consolas'
                    term_run.font.color.rgb = RGBColor(80, 80, 80)

            # 錯誤訊息
            error = tc.get('error', '')
            if error:
                error_p = doc.add_paragraph()
                error_run = error_p.add_run(f'Error: {error[:300]}')
                error_run.font.size = Pt(9)
                error_run.font.color.rgb = RGBColor(244, 67, 54)

            doc.add_paragraph()  # 間隔

    # ========== 各類型長條圖 ==========
    if include_charts and HAS_MATPLOTLIB and tests:
        doc.add_heading('測試結果分布', level=2)

        chart_data = create_test_type_chart(tests)
        if chart_data:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp.write(chart_data)
                tmp_path = tmp.name

            doc.add_picture(tmp_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            Path(tmp_path).unlink(missing_ok=True)

    doc.add_page_break()

    # ========== 測試類型說明 ==========
    doc.add_heading('測試類型說明', level=1)

    descriptions = [
        ('UIT (Unit Integration Testing)',
         '單元測試驗證各個模組、函數的正確性。使用 Vitest/Jest 框架執行，並產生程式碼覆蓋率報告。'),
        ('Smoke Test (煙霧測試)',
         '快速驗證系統關鍵路徑是否正常運作。包含應用程式啟動、頁面載入、API 健康檢查等基本功能。'),
        ('E2E (End-to-End Testing)',
         '端對端測試模擬真實使用情境，驗證完整的使用者流程。使用 Playwright 自動化測試框架執行。'),
        ('UAT (User Acceptance Testing)',
         '使用者驗收測試從業務角度驗證系統符合需求規格。測試案例依據使用者角色設計，確保系統滿足業務需求。'),
    ]

    for title, desc in descriptions:
        p = doc.add_paragraph()
        title_run = p.add_run(title + ': ')
        title_run.bold = True
        p.add_run(desc)
        doc.add_paragraph()

    # ========== 額外截圖 (可選，用於補充說明) ==========
    if screenshots:
        doc.add_page_break()
        doc.add_heading('補充截圖', level=1)

        for i, screenshot_path in enumerate(screenshots):
            if Path(screenshot_path).exists():
                p = doc.add_paragraph()
                run = p.add_run(f'截圖 {i + 1}')
                run.bold = True
                run.font.size = Pt(12)

                doc.add_picture(screenshot_path, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()

    # ========== 頁尾 ==========
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Generated by DashAI DevTools')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)

    # 儲存文件
    output_file = Path(output_path)
    doc.save(output_file)

    return str(output_file)


def take_screenshots(project_path: str, urls: List[str] = None) -> List[str]:
    """
    使用 Puppeteer 截取頁面截圖

    Args:
        project_path: 專案路徑
        urls: 要截圖的 URL 列表

    Returns:
        截圖檔案路徑列表
    """
    import subprocess
    import tempfile

    if not urls:
        # 預設截圖頁面
        urls = [
            'https://smai-mes.vercel.app/',
            'https://smai-mes.vercel.app/warehouse/transfer',
            'https://smai-mes.vercel.app/warehouse/line-side',
        ]

    screenshots = []
    screenshot_dir = Path(tempfile.mkdtemp(prefix='test-screenshots-'))

    # Puppeteer 截圖腳本
    for i, url in enumerate(urls):
        screenshot_path = screenshot_dir / f'screenshot-{i+1}.png'
        script = f'''
const puppeteer = require("puppeteer");
(async () => {{
  const browser = await puppeteer.launch({{ headless: "new" }});
  const page = await browser.newPage();
  await page.setViewport({{ width: 1920, height: 1080 }});
  await page.goto("{url}", {{ waitUntil: "networkidle0", timeout: 30000 }});
  await new Promise(r => setTimeout(r, 2000));
  await page.screenshot({{ path: "{screenshot_path}", fullPage: false }});
  await browser.close();
}})();
'''
        try:
            # 找有 puppeteer 的目錄
            puppeteer_dirs = [
                '/Users/dash/Documents/github/smai-process-vision',
                '/Users/dash/Documents/github/MES',
                project_path
            ]

            for pdir in puppeteer_dirs:
                if (Path(pdir) / 'node_modules' / 'puppeteer').exists():
                    result = subprocess.run(
                        ['node', '-e', script],
                        cwd=pdir,
                        capture_output=True,
                        timeout=60
                    )
                    if result.returncode == 0 and screenshot_path.exists():
                        screenshots.append(str(screenshot_path))
                        console.print(f"[dim]  截圖: {url}[/dim]")
                    break
        except Exception as e:
            console.print(f"[yellow]截圖失敗: {url} - {e}[/yellow]")

    return screenshots


def run_and_generate_report(
    project_path: str,
    output_path: str = None,
    test_types: List[str] = None,
    include_screenshots: bool = True,
    screenshot_urls: List[str] = None
) -> Dict:
    """
    執行測試並生成 Word 報告

    Args:
        project_path: 專案路徑
        output_path: 輸出路徑 (預設為專案目錄下的 test-report.docx)
        test_types: 測試類型列表
        include_screenshots: 是否包含截圖 (預設 True)
        screenshot_urls: 要截圖的 URL 列表

    Returns:
        結果字典
    """
    from .test_suite import TestSuiteRunner

    project = Path(project_path)
    project_name = project.name

    if not output_path:
        output_path = str(project / f'{project_name}-test-report.docx')

    # 執行測試
    console.print(f"[cyan]執行 {project_name} 測試套件...[/cyan]")
    runner = TestSuiteRunner(project_path)
    suite_result = runner.run_all(test_types)

    # 準備測試結果 (包含測試案例明細)
    test_results = {
        'project': project_name,
        'timestamp': suite_result.timestamp,
        'overall_success': suite_result.overall_success,
        'summary': {
            'total_passed': suite_result.total_passed,
            'total_failed': suite_result.total_failed,
            'total_duration': suite_result.total_duration,
            'coverage': suite_result.coverage
        },
        'tests': {
            k: {
                'success': v.success,
                'passed': v.passed,
                'failed': v.failed,
                'duration': v.duration,
                'coverage': v.coverage,
                'test_cases': [
                    {
                        'name': tc.name,
                        'status': tc.status,
                        'duration': tc.duration,
                        'error': tc.error,
                        'screenshot': tc.screenshot,
                        'api_response': tc.api_response,
                        'terminal_output': tc.terminal_output
                    }
                    for tc in v.test_cases
                ]
            }
            for k, v in suite_result.results.items()
        }
    }

    # 額外截圖 (Playwright 測試已自動截圖，這裡只用於補充)
    screenshots = []
    if include_screenshots and screenshot_urls:
        console.print(f"[cyan]擷取補充截圖...[/cyan]")
        screenshots = take_screenshots(project_path, screenshot_urls)

    # 生成報告
    console.print(f"[cyan]生成 Word 報告...[/cyan]")
    report_path = generate_word_report(
        project_name=project_name,
        test_results=test_results,
        output_path=output_path,
        screenshots=screenshots,
        include_charts=True
    )

    console.print(f"[green]報告已生成: {report_path}[/green]")

    return {
        'success': True,
        'report_path': report_path,
        'test_results': test_results
    }
